"""
Tests for the /api/analyze router and health endpoint.

Tests endpoint validation, SSE streaming structure, error handling,
and the health check response using FastAPI TestClient.
"""

from __future__ import annotations

import io
import json
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from fastapi.testclient import TestClient

from backend.main import app
from backend.models.schemas import (
    AggregateScoreBreakdown,
    AnalysisReport,
    ClauseCategory,
    RiskTier,
    Severity,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_minimal_docx() -> bytes:
    """Create a minimal valid DOCX for upload testing."""
    from docx import Document
    doc = Document()
    doc.add_paragraph("Section 1. Intellectual Property Rights")
    doc.add_paragraph(
        "The Company retains all rights to any work product created "
        "during the term of employment, including but not limited to "
        "inventions, designs, software, and written materials produced "
        "in the course of performing assigned duties."
    )
    doc.add_paragraph("Section 2. Non-Compete Agreement")
    doc.add_paragraph(
        "The Employee agrees not to engage in any competing business "
        "within a radius of 50 miles for a period of 24 months following "
        "termination of employment for any reason whatsoever."
    )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_mock_report() -> AnalysisReport:
    """Create a minimal mock AnalysisReport."""
    return AnalysisReport(
        document_type="Employment Contract",
        total_clauses=2,
        flagged_clauses=1,
        risk_tier=RiskTier.MODERATE,
        aggregate_risk_score=35.0,
        score_breakdown=AggregateScoreBreakdown(
            total_clause_score=5.0,
            max_possible_score=50.0,
            raw_percentage=35.0,
            clause_count=2,
        ),
        clause_reports=[],
        category_heatmap=[],
        contradictions=[],
    )


# ---------------------------------------------------------------------------
# Health endpoint
# ---------------------------------------------------------------------------

class TestHealthEndpoint:
    """Tests for the /health endpoint."""

    def test_health_returns_200(self):
        """Health endpoint should return 200 OK."""
        client = TestClient(app)
        response = client.get("/health")
        assert response.status_code == 200

    def test_health_returns_ok_status(self):
        """Health endpoint should return status: ok."""
        client = TestClient(app)
        response = client.get("/health")
        data = response.json()
        assert data["status"] == "ok"

    def test_health_returns_version(self):
        """Health endpoint should return a version string."""
        client = TestClient(app)
        response = client.get("/health")
        data = response.json()
        assert "version" in data
        assert data["version"] == "1.0.0"


# ---------------------------------------------------------------------------
# /api/analyze — input validation
# ---------------------------------------------------------------------------

class TestAnalyzeEndpointValidation:
    """Tests for upload validation at the /api/analyze endpoint."""

    def test_missing_file_returns_422(self):
        """Request without a file should return 422 Unprocessable Entity."""
        client = TestClient(app)
        response = client.post("/api/analyze")
        assert response.status_code == 422

    def test_empty_file_returns_400(self):
        """Empty file upload should return 400 Bad Request."""
        client = TestClient(app)
        response = client.post(
            "/api/analyze",
            files={"file": ("empty.pdf", b"", "application/pdf")},
        )
        assert response.status_code == 400

    def test_unsupported_file_type_returns_400(self):
        """Unsupported file type should return 400 Bad Request."""
        client = TestClient(app)
        response = client.post(
            "/api/analyze",
            files={"file": ("contract.txt", b"some text content here", "text/plain")},
        )
        assert response.status_code == 400

    def test_oversized_file_returns_400(self):
        """File exceeding 10MB should return 400 Bad Request."""
        from backend.utils.validators import MAX_FILE_SIZE_BYTES
        client = TestClient(app)
        large_content = b"%PDF-1.4" + b"x" * (MAX_FILE_SIZE_BYTES + 1)
        response = client.post(
            "/api/analyze",
            files={"file": ("large.pdf", large_content, "application/pdf")},
        )
        assert response.status_code == 400

    def test_mismatched_extension_returns_400(self):
        """File with PDF extension but DOCX content should return 400."""
        client = TestClient(app)
        docx_bytes = _make_minimal_docx()
        # DOCX bytes (starts with PK\x03\x04) but .pdf extension — should fail cross-check
        response = client.post(
            "/api/analyze",
            files={"file": ("contract.pdf", docx_bytes, "application/octet-stream")},
        )
        assert response.status_code == 400


# ---------------------------------------------------------------------------
# /api/analyze — SSE streaming
# ---------------------------------------------------------------------------

class TestAnalyzeEndpointStreaming:
    """Tests for SSE streaming from /api/analyze."""

    @patch("backend.routers.analyze.parse_document")
    @patch("backend.routers.analyze.classify_clauses")
    @patch("backend.routers.analyze.batch_embed_clauses")
    @patch("backend.routers.analyze._process_single_clause")
    @patch("backend.routers.analyze.detect_contradictions")
    @patch("backend.routers.analyze.generate_report")
    @patch("backend.routers.analyze.upload_document")
    @patch("backend.routers.analyze.delete_document")
    def test_valid_docx_returns_streaming_response(
        self,
        mock_delete: AsyncMock,
        mock_upload: AsyncMock,
        mock_report: MagicMock,
        mock_contradictions: AsyncMock,
        mock_process: AsyncMock,
        mock_embed: AsyncMock,
        mock_classify: AsyncMock,
        mock_parse: MagicMock,
    ):
        """Valid DOCX upload should return a streaming response."""
        from backend.models.schemas import Clause, ClauseReport
        mock_parse.return_value = [
            Clause(id="clause_001", text="The employee agrees to all terms.", section="General")
        ]
        mock_classify.return_value = mock_parse.return_value
        mock_embed.return_value = {}
        mock_process.return_value = ClauseReport(
            clause=mock_parse.return_value[0],
            category=ClauseCategory.IP_TRANSFER,
            severity=Severity.HIGH,
        )
        mock_contradictions.return_value = []
        mock_report.return_value = _make_mock_report()
        mock_upload.return_value = None
        mock_delete.return_value = None

        client = TestClient(app)
        docx_bytes = _make_minimal_docx()
        response = client.post(
            "/api/analyze",
            files={"file": ("contract.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert response.status_code == 200
        assert "text/event-stream" in response.headers.get("content-type", "")

    @patch("backend.routers.analyze.parse_document")
    @patch("backend.routers.analyze.upload_document")
    def test_parse_failure_streams_error_event(
        self,
        mock_upload: AsyncMock,
        mock_parse: MagicMock,
    ):
        """Parse failure should stream an error SSE event."""
        mock_upload.return_value = None
        mock_parse.side_effect = ValueError("PDF contains no extractable text")

        client = TestClient(app)
        docx_bytes = _make_minimal_docx()
        response = client.post(
            "/api/analyze",
            files={"file": ("contract.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert response.status_code == 200
        content = response.text
        # Should contain an error event
        assert "error" in content.lower()


# ---------------------------------------------------------------------------
# Security headers
# ---------------------------------------------------------------------------

class TestSecurityHeaders:
    """Tests for security headers middleware."""

    def test_x_content_type_options_header_present(self):
        """X-Content-Type-Options header should be set."""
        client = TestClient(app)
        response = client.get("/health")
        assert response.headers.get("x-content-type-options") == "nosniff"

    def test_x_frame_options_header_present(self):
        """X-Frame-Options header should be set."""
        client = TestClient(app)
        response = client.get("/health")
        assert response.headers.get("x-frame-options") == "DENY"

    def test_x_xss_protection_header_present(self):
        """X-XSS-Protection header should be set."""
        client = TestClient(app)
        response = client.get("/health")
        assert response.headers.get("x-xss-protection") == "1; mode=block"

    def test_referrer_policy_header_present(self):
        """Referrer-Policy header should be set."""
        client = TestClient(app)
        response = client.get("/health")
        assert response.headers.get("referrer-policy") == "strict-origin-when-cross-origin"
