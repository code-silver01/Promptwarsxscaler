"""
Pytest configuration and shared fixtures for LexGuard One tests.

Provides comprehensive test fixtures, mock configurations, and test utilities
for all components of the contract analysis platform.
"""

import asyncio
import io
import json
import os
import tempfile
from datetime import datetime
from typing import Any, Dict, List, Optional
from unittest.mock import AsyncMock, MagicMock

import pytest
from docx import Document
from fastapi.testclient import TestClient

from backend.models.schemas import (
    AnalysisReport,
    BenchmarkComparison,
    Clause,
    ClauseCategory,
    ClauseReport,
    ConsequenceChain,
    DefenseAgentOutput,
    NegotiationSuggestion,
    RiskAgentOutput,
    RiskTier,
    Severity,
    VerdictAgentOutput,
)


# Test Configuration
@pytest.fixture(scope="session")
def event_loop():
    """Create an instance of the default event loop for the test session."""
    loop = asyncio.get_event_loop_policy().new_event_loop()
    yield loop
    loop.close()


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield tmpdir


@pytest.fixture
def mock_env_vars(monkeypatch):
    """Mock environment variables for testing."""
    test_env = {
        "GOOGLE_CLOUD_PROJECT": "test-project",
        "FIRESTORE_COLLECTION": "test-contracts",
        "GCS_BUCKET": "test-bucket",
        "GEMINI_API_KEY": "test-key",
        "SECRET_MANAGER_PROJECT": "test-project",
    }
    for key, value in test_env.items():
        monkeypatch.setenv(key, value)
    return test_env


# Document Fixtures
@pytest.fixture
def sample_pdf_bytes():
    """Provide minimal PDF bytes for testing."""
    return b"%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj 2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj 3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]>>endobj xref 0 4 0000000000 65535 f 0000000010 00000 n 0000000053 00000 n 0000000125 00000 n trailer<</Size 4/Root 1 0 R>>startxref 185 %%EOF"


@pytest.fixture
def sample_docx_bytes():
    """Provide valid DOCX bytes for testing."""
    doc = Document()
    doc.add_paragraph("Section 1. Employment Terms")
    doc.add_paragraph(
        "The Employee agrees to perform all duties as assigned by "
        "the Company including but not limited to software development, "
        "testing, and documentation tasks as required by management."
    )
    doc.add_paragraph("Section 2. Intellectual Property")
    doc.add_paragraph(
        "All work product, inventions, and intellectual property "
        "created by the Employee during the term of employment shall "
        "be the sole and exclusive property of the Company."
    )
    doc.add_paragraph("Section 3. Non-Compete Agreement")
    doc.add_paragraph(
        "Employee shall not engage in any business that competes with "
        "the Company for a period of 2 years after termination."
    )
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def complex_contract_docx():
    """Create a complex contract with multiple risky clauses."""
    doc = Document()
    
    # Title
    doc.add_heading("Software Development Agreement", 0)
    
    # IP Transfer clause
    doc.add_paragraph("1. Intellectual Property Rights")
    doc.add_paragraph(
        "All work product, including but not limited to software code, "
        "documentation, inventions, discoveries, improvements, and any "
        "intellectual property created by Developer during the term of "
        "this agreement, whether during business hours or not, and "
        "whether using Company resources or not, shall become the "
        "exclusive property of the Company."
    )
    
    # Liability limitation
    doc.add_paragraph("2. Limitation of Liability")
    doc.add_paragraph(
        "In no event shall Company be liable for any indirect, "
        "incidental, special, consequential, or punitive damages, "
        "including but not limited to loss of profits, data, or use, "
        "regardless of the theory of liability."
    )
    
    # Auto-renewal
    doc.add_paragraph("3. Term and Renewal")
    doc.add_paragraph(
        "This agreement shall automatically renew for successive "
        "one-year periods unless either party provides written notice "
        "of non-renewal at least 90 days prior to the expiration date."
    )
    
    # Arbitration
    doc.add_paragraph("4. Dispute Resolution")
    doc.add_paragraph(
        "Any dispute arising under this agreement shall be resolved "
        "exclusively through binding arbitration in accordance with "
        "the rules of the American Arbitration Association. The "
        "prevailing party shall be entitled to attorney's fees."
    )
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# Model Fixtures
@pytest.fixture
def sample_clause() -> Clause:
    """Create a sample clause for testing."""
    return Clause(
        id="clause_001",
        text=(
            "All work product, inventions, and intellectual property "
            "created by the Employee during the term of employment shall "
            "be the sole and exclusive property of the Company."
        ),
        section="Intellectual Property",
        category=ClauseCategory.IP_TRANSFER,
        raw_span=[100, 250],
    )


@pytest.fixture
def high_risk_clause() -> Clause:
    """Create a high-risk clause for testing."""
    return Clause(
        id="clause_high_risk",
        text=(
            "Employee agrees that all work product, including personal "
            "projects developed outside of work hours using any equipment "
            "or knowledge gained during employment, shall be owned by Company. "
            "This includes any inventions, software, or intellectual property "
            "created up to 2 years after termination of employment."
        ),
        section="Intellectual Property",
        category=ClauseCategory.IP_TRANSFER,
    )


@pytest.fixture
def sample_clauses() -> List[Clause]:
    """Create a list of sample clauses for testing."""
    return [
        Clause(
            id="clause_001",
            text="Standard employment terms and conditions apply.",
            section="General",
            category=None,
        ),
        Clause(
            id="clause_002",
            text=(
                "All intellectual property created during employment "
                "belongs to the Company."
            ),
            section="IP Rights",
            category=ClauseCategory.IP_TRANSFER,
        ),
        Clause(
            id="clause_003",
            text=(
                "Employee shall not compete with Company for 1 year "
                "after termination."
            ),
            section="Non-Compete",
            category=ClauseCategory.NON_COMPETE,
        ),
    ]


# Agent Output Fixtures
@pytest.fixture
def mock_risk_output() -> RiskAgentOutput:
    """Create mock Risk Agent output."""
    return RiskAgentOutput(
        risk_position=(
            "This clause grants the company extremely broad ownership "
            "rights over all employee work product, including personal projects."
        ),
        key_phrases=[
            "all work product",
            "personal projects",
            "outside of work hours",
            "2 years after termination",
        ],
        worst_case=(
            "Company could claim ownership of employee's side projects, "
            "portfolio work, and future innovations for 2 years post-employment."
        ),
        reasoning=(
            "Step 1: Identified overly broad scope including 'personal projects'. "
            "Step 2: Noted temporal extension beyond employment period. "
            "Step 3: Assessed enforceability and potential for abuse."
        ),
    )


@pytest.fixture
def mock_defense_output() -> DefenseAgentOutput:
    """Create mock Defense Agent output."""
    return DefenseAgentOutput(
        defense_position=(
            "The clause may be limited by state laws protecting employee "
            "inventions and could be interpreted narrowly to work-related IP only."
        ),
        favorable_phrases=[
            "during employment",
            "using company equipment",
            "work-related knowledge",
        ],
        best_case=(
            "Courts may limit enforcement to work directly related to "
            "company business and created using company resources."
        ),
        reasoning=(
            "Step 1: Identified potential state law protections. "
            "Step 2: Analyzed limiting language in the clause. "
            "Step 3: Considered judicial precedents for narrow interpretation."
        ),
    )


@pytest.fixture
def mock_verdict_output() -> VerdictAgentOutput:
    """Create mock Verdict Agent output."""
    return VerdictAgentOutput(
        verdict=(
            "This clause presents significant risk due to its broad scope "
            "but may be limited by state employment laws."
        ),
        severity=Severity.HIGH,
        confidence=0.85,
        risk_category="IP_TRANSFER",
        plain_english=(
            "This clause could let your employer claim ownership of your "
            "personal coding projects and side businesses. While some states "
            "have laws protecting employees, the risk is still high."
        ),
        reasoning=(
            "Step 1: Weighed broad scope against legal protections. "
            "Step 2: Considered enforceability in different jurisdictions. "
            "Step 3: Assessed real-world impact on employee rights."
        ),
    )


@pytest.fixture
def mock_consequence_chain() -> ConsequenceChain:
    """Create mock consequence chain."""
    return ConsequenceChain(
        trigger_condition="Employee creates personal software project",
        immediate_consequence="Company claims ownership of the project",
        downstream_impact=(
            "Employee loses rights to personal work, potential legal disputes, "
            "chilling effect on innovation"
        ),
        worst_case_scenario=(
            "Employee's successful startup or app gets claimed by former employer, "
            "resulting in loss of business and potential millions in damages"
        ),
    )


@pytest.fixture
def mock_negotiation_suggestion() -> NegotiationSuggestion:
    """Create mock negotiation suggestion."""
    return NegotiationSuggestion(
        original_clause_text=(
            "All work product created during employment belongs to Company."
        ),
        suggested_alternative_text=(
            "Work product created during business hours using Company resources "
            "and directly related to Company business belongs to Company. "
            "Personal projects created outside work hours without Company "
            "resources remain employee property."
        ),
        why_safer=(
            "This revision limits Company ownership to work-related IP and "
            "protects employee rights to personal projects and innovations."
        ),
    )


@pytest.fixture
def mock_benchmark_comparison() -> BenchmarkComparison:
    """Create mock benchmark comparison."""
    return BenchmarkComparison(
        percentile=85.0,
        summary=(
            "This clause is more restrictive than 85% of similar IP "
            "assignment clauses in our benchmark dataset."
        ),
        top_matches=[
            "All inventions and work product shall be Company property",
            "Employee assigns all rights to work-related intellectual property",
            "Company owns all IP created during employment period",
        ],
    )


@pytest.fixture
def sample_clause_report(
    sample_clause: Clause,
    mock_risk_output: RiskAgentOutput,
    mock_defense_output: DefenseAgentOutput,
    mock_verdict_output: VerdictAgentOutput,
) -> ClauseReport:
    """Create a complete clause report for testing."""
    return ClauseReport(
        clause=sample_clause,
        category=ClauseCategory.IP_TRANSFER,
        severity=Severity.HIGH,
        confidence=0.85,
        risk_position=mock_risk_output,
        defense_position=mock_defense_output,
        verdict=mock_verdict_output,
        plain_english=mock_verdict_output.plain_english,
    )


@pytest.fixture
def sample_analysis_report() -> AnalysisReport:
    """Create a sample analysis report."""
    return AnalysisReport(
        document_type="Employment Agreement",
        total_clauses=15,
        flagged_clauses=5,
        risk_tier=RiskTier.HIGH,
        aggregate_score=72.5,
        processing_time_ms=15000,
        clause_reports=[],
        score_breakdown=None,
        category_heatmap=[],
        contradictions=[],
        executive_summary="High-risk contract with multiple concerning clauses.",
        key_recommendations=[
            "Negotiate IP assignment clause",
            "Add termination protections",
            "Clarify liability limitations",
        ],
    )


# Mock Fixtures
@pytest.fixture
def mock_gemini_client():
    """Mock Gemini API client."""
    mock = AsyncMock()
    mock.call_gemini.return_value = {
        "risk_position": "Test risk position",
        "key_phrases": ["test phrase"],
        "worst_case": "Test worst case",
        "reasoning": "Test reasoning",
    }
    return mock


@pytest.fixture
def mock_firestore_client():
    """Mock Firestore client."""
    mock = MagicMock()
    mock.collection.return_value.document.return_value.set = AsyncMock()
    mock.collection.return_value.document.return_value.get = AsyncMock()
    return mock


@pytest.fixture
def mock_gcs_client():
    """Mock Google Cloud Storage client."""
    mock = MagicMock()
    mock.upload_document = AsyncMock(return_value="gs://test-bucket/test-file.pdf")
    mock.delete_document = AsyncMock(return_value=True)
    return mock


@pytest.fixture
def mock_vertex_ai():
    """Mock Vertex AI client."""
    mock = MagicMock()
    mock.embed_texts = AsyncMock(return_value=[[0.1, 0.2, 0.3] * 384])
    return mock


# Test Client Fixtures
@pytest.fixture
def test_client():
    """Create FastAPI test client."""
    from backend.main import app
    return TestClient(app)


@pytest.fixture
def authenticated_client(test_client):
    """Create authenticated test client."""
    # Add authentication headers if needed
    test_client.headers.update({"Authorization": "Bearer test-token"})
    return test_client


# Utility Fixtures
@pytest.fixture
def signer_profile() -> Dict[str, Any]:
    """Create a sample signer profile."""
    return {
        "name": "John Doe",
        "occupation": "Software Developer",
        "experience": "Senior",
        "company_size": "Startup",
        "industry": "Technology",
        "location": "California",
        "risk_tolerance": "Low",
    }


@pytest.fixture
def performance_benchmark():
    """Performance benchmarking fixture."""
    import time
    
    class PerformanceBenchmark:
        def __init__(self):
            self.start_time = None
            self.end_time = None
        
        def start(self):
            self.start_time = time.time()
        
        def stop(self):
            self.end_time = time.time()
        
        @property
        def duration(self):
            if self.start_time and self.end_time:
                return self.end_time - self.start_time
            return None
    
    return PerformanceBenchmark()


# Database Fixtures
@pytest.fixture
async def clean_database():
    """Clean test database before and after tests."""
    # Setup: Clean database
    yield
    # Teardown: Clean database


# Security Test Fixtures
@pytest.fixture
def malicious_pdf_bytes():
    """Create malicious PDF bytes for security testing."""
    return b"%PDF-1.4\n<script>alert('xss')</script>\n%%EOF"


@pytest.fixture
def oversized_file_bytes():
    """Create oversized file for testing file size limits."""
    return b"x" * (50 * 1024 * 1024)  # 50MB file


# Parametrized Fixtures
@pytest.fixture(params=[
    ClauseCategory.IP_TRANSFER,
    ClauseCategory.NON_COMPETE,
    ClauseCategory.ARBITRATION,
    ClauseCategory.LIABILITY_LIMITATION,
])
def risky_clause_category(request):
    """Parametrized fixture for risky clause categories."""
    return request.param


@pytest.fixture(params=[Severity.LOW, Severity.MEDIUM, Severity.HIGH])
def severity_level(request):
    """Parametrized fixture for severity levels."""
    return request.param
